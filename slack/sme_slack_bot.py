#!/usr/bin/env python3
"""
Enhanced SME Slack Bot - Auto-Monitor Mode

Automatically captures and processes ALL messages in monitored channel:
- Text messages
- Web links (auto-scrapes content)
- PDF files
- Word documents
- NO @mention required!

Integrates with existing expert_examples database.
"""

import os
import re
import sys
import hashlib
import tempfile
import sqlite3
import warnings
from pathlib import Path
from datetime import datetime

# Suppress urllib3 NotOpenSSLWarning on macOS
warnings.filterwarnings('ignore', message='.*urllib3 v2 only supports OpenSSL.*')

try:
    from slack_bolt import App
    from slack_bolt.adapter.socket_mode import SocketModeHandler
except ImportError:
    print("❌ slack-bolt not installed")
    print("\nInstall it with:")
    print("  pip3 install slack-bolt slack-sdk")
    sys.exit(1)

try:
    import requests
    from bs4 import BeautifulSoup
except ImportError:
    print("❌ Web scraping libraries not installed")
    print("\nInstall them with:")
    print("  pip3 install requests beautifulsoup4")
    sys.exit(1)

# Add project root to path for imports
# We're in slack/ folder, so go up one level to project root
project_root = Path(__file__).parent.parent.resolve()
sys.path.insert(0, str(project_root))
from scripts.add_expert_example import add_expert_example

# Configuration from environment variables
SLACK_BOT_TOKEN = os.environ.get("SLACK_BOT_TOKEN")
SLACK_APP_TOKEN = os.environ.get("SLACK_APP_TOKEN")
SME_CHANNEL_NAME = os.environ.get("SME_CHANNEL", "sme-knowledge")
# Database path - default to grants.db in project root (one level up from slack/)
DB_PATH = os.environ.get("DB_PATH", str(project_root / "grants.db"))

if not SLACK_BOT_TOKEN or not SLACK_APP_TOKEN:
    print("❌ Missing Slack tokens!")
    print("\nSet environment variables:")
    print("  export SLACK_BOT_TOKEN=xoxb-your-bot-token")
    print("  export SLACK_APP_TOKEN=xapp-your-app-token")
    print("\nOr create a .env file (see .env.example)")
    sys.exit(1)

# Initialize Slack app
app = App(token=SLACK_BOT_TOKEN)

# Message deduplication cache
processed_messages = set()


# ==================== MESSAGE DEDUPLICATION ====================

def should_process_message(event):
    """
    Determine if we should process this message.
    Prevents duplicate processing and filters out bot messages.
    """
    # Generate unique message ID
    msg_id = f"{event.get('channel')}_{event.get('ts')}"

    # Skip if already processed
    if msg_id in processed_messages:
        return False

    # Skip bot messages (including our own)
    if event.get('bot_id') or event.get('subtype') == 'bot_message':
        return False

    # Skip thread broadcast messages (they're duplicates)
    if event.get('subtype') == 'thread_broadcast':
        return False

    # Skip message edits, deletes, etc.
    if event.get('subtype') and event.get('subtype') not in [None, 'file_share']:
        return False

    # Add to processed set
    processed_messages.add(msg_id)

    # Keep cache size reasonable (max 1000 messages)
    if len(processed_messages) > 1000:
        # Remove oldest half
        to_remove = list(processed_messages)[:500]
        for msg in to_remove:
            processed_messages.discard(msg)

    return True


# ==================== WEB SCRAPING ====================

def scrape_web_content(url: str) -> dict:
    """
    Scrape content from a web URL.

    Args:
        url: URL to scrape

    Returns:
        Dict with title, content, and metadata
    """
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36'
        }
        response = requests.get(url, headers=headers, timeout=15)
        response.raise_for_status()

        soup = BeautifulSoup(response.text, 'html.parser')

        # Remove unwanted elements
        for element in soup(['script', 'style', 'nav', 'header', 'footer', 'aside']):
            element.decompose()

        # Get title
        title = soup.find('title')
        title_text = title.string if title else url

        # Try to find main content
        content_selectors = [
            'main', 'article', '[role="main"]',
            '.content', '#content', '.post-content'
        ]

        main_content = None
        for selector in content_selectors:
            main_content = soup.select_one(selector)
            if main_content:
                break

        if not main_content:
            main_content = soup.find('body') or soup

        # Extract text
        text_parts = []
        for elem in main_content.find_all(['h1', 'h2', 'h3', 'p', 'li']):
            text = elem.get_text(strip=True)
            if len(text) > 20:  # Skip short fragments
                text_parts.append(text)

        content = '\n\n'.join(text_parts)

        # Extract key info for grant pages
        deadline = None
        funding = None

        if any(domain in url for domain in ['innovateuk', 'nihr', 'ukri']):
            deadline_elem = soup.find(text=re.compile(r'deadline|closes', re.I))
            if deadline_elem:
                deadline = deadline_elem.strip()

            funding_elem = soup.find(text=re.compile(r'£[\d,]+', re.I))
            if funding_elem:
                funding = funding_elem.strip()

        return {
            'title': title_text,
            'content': content[:5000],  # Limit to 5000 chars
            'url': url,
            'deadline': deadline,
            'funding': funding,
            'success': True
        }

    except Exception as e:
        return {
            'title': url,
            'content': '',
            'url': url,
            'success': False,
            'error': str(e)
        }


def extract_pdf_text(file_path: str) -> str:
    """
    Extract text from PDF file.

    Args:
        file_path: Path to PDF file

    Returns:
        Extracted text
    """
    try:
        import PyPDF2

        text = ""
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

        return text
    except Exception as e:
        print(f"PDF extraction failed: {e}")
        return ""


def extract_docx_text(file_path: str) -> str:
    """
    Extract text from Word document.

    Args:
        file_path: Path to .docx file

    Returns:
        Extracted text
    """
    try:
        from docx import Document

        doc = Document(file_path)
        text_parts = []

        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)

        return '\n\n'.join(text_parts)
    except Exception as e:
        print(f"DOCX extraction failed: {e}")
        return ""


# ==================== CONTENT DETECTION ====================

def detect_category(text: str) -> str:
    """Auto-detect message category based on content."""
    text_lower = text.lower()

    if any(word in text_lower for word in ['eligibility', 'eligible', 'criteria', 'requirements']):
        return 'eligibility'

    if any(word in text_lower for word in ['application', 'apply', 'deadline', 'submit']):
        return 'application_process'

    if any(word in text_lower for word in ['strategy', 'position', 'approach', 'angle', 'frame']):
        return 'positioning'

    if any(word in text_lower for word in ['feasibility', 'proof of concept', 'poc']):
        return 'feasibility'

    if any(word in text_lower for word in ['loan', 'credit', 'repay', 'financing']):
        return 'financing'

    if any(word in text_lower for word in ['biomedical catalyst', 'i4i', 'smart grant', 'ktp', 'sbri']):
        return 'grant_explanation'

    return 'strategy'


def extract_grant_name(text: str) -> str:
    """Extract grant name from message."""
    grant_patterns = [
        r'(Biomedical Catalyst)',
        r'(Invention for Innovation|i4i)',
        r'(Smart Grant)',
        r'(Knowledge Transfer Partnership|KTP)',
        r'(SBRI)',
        r'(Innovation Loan)',
        r'(Launchpad)',
        r'(NIHR [A-Z]{2,4})',
        r'(Innovate UK [A-Za-z\s]+)',
    ]

    for pattern in grant_patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return match.group(1)

    return None


def assess_quality(text: str, source: str) -> int:
    """Assess quality score for content."""
    score = 3  # Base score

    # Length bonus
    if len(text) > 500:
        score += 1

    # Source bonuses
    if source in ['pdf', 'docx']:
        score += 1  # Formal documents
    elif source == 'url':
        score += 0.5

    # Specific numbers/amounts
    if re.search(r'£[\d,]+[KkMm]?', text):
        score += 0.5
    if re.search(r'\d+%', text):
        score += 0.5

    # Strategic keywords
    strategic_words = ['strategy', 'position', 'angle', 'approach', 'recommend', 'suggest']
    if any(word in text.lower() for word in strategic_words):
        score += 0.5

    # Examples
    if any(phrase in text.lower() for phrase in ['for example', 'in my experience', 'typically']):
        score += 0.5

    return min(5, round(score))


def generate_query(text: str, grant_name: str = None, source: str = 'slack') -> str:
    """Generate natural user query."""
    text_lower = text.lower()

    # Source-specific queries
    if source == 'url':
        if grant_name:
            return f"Tell me about {grant_name}"
        return "What are the current grant opportunities?"

    if source in ['pdf', 'docx']:
        if grant_name:
            return f"What are the guidelines for {grant_name}?"
        return "What are the application guidelines?"

    # Content-based queries
    if 'strategy' in text_lower or 'approach' in text_lower:
        if grant_name:
            return f"What's the best strategy for applying to {grant_name}?"
        return "What's the best approach for grant applications?"

    if 'eligibility' in text_lower or 'criteria' in text_lower:
        if grant_name:
            return f"Am I eligible for {grant_name}?"
        return "What are the eligibility requirements?"

    if 'application' in text_lower or 'deadline' in text_lower:
        if grant_name:
            return f"How do I apply for {grant_name}?"
        return "What's the application process?"

    if grant_name:
        return f"Tell me about {grant_name}"

    return "What grant funding is available?"


# ==================== SLACK EVENT HANDLERS ====================

@app.event("message")
def handle_message(event, say, client, logger):
    """
    AUTO-MONITOR: Handle ALL incoming Slack messages (no @mention required).
    Processes text, URLs, PDFs, and Word docs automatically.
    """
    # Check if we should process this message (deduplication + bot filtering)
    if not should_process_message(event):
        return

    # Only process messages in the SME channel
    channel_id = event.get('channel')

    try:
        channel_info = client.conversations_info(channel=channel_id)
        channel_name = channel_info['channel']['name']

        if channel_name != SME_CHANNEL_NAME:
            return
    except Exception as e:
        logger.error(f"Failed to get channel info: {e}")
        return

    # Get message content
    text = event.get('text', '').strip()
    user_id = event.get('user')

    # Remove bot mentions from text (clean up @bot mentions if present)
    # This handles cases where users still mention the bot
    text = re.sub(r'<@[A-Z0-9]+>', '', text).strip()

    # Log what we're processing (for visibility)
    logger.info(f"🎯 Auto-processing in #{channel_name}: {text[:80]}...")

    # Get user info
    try:
        user_info = client.users_info(user=user_id)
        user_name = user_info['user']['real_name'] or user_info['user']['name']
    except:
        user_name = "SME"

    imported_items = []
    all_topics = set()

    # ===== PROCESS WEB LINKS =====
    urls = re.findall(r'https?://[^\s<>"\{\}|\\^`\[\]]+', text)
    for url in urls:
        logger.info(f"📎 Processing URL: {url}")

        scraped = scrape_web_content(url)

        if scraped['success'] and scraped['content']:
            # Combine URL content with user's message
            full_content = f"# {scraped['title']}\n\nSource: {url}\n\n{scraped['content']}"

            if scraped['deadline']:
                full_content += f"\n\nDeadline: {scraped['deadline']}"
            if scraped['funding']:
                full_content += f"\n\nFunding: {scraped['funding']}"

            # Add user's context
            clean_text = re.sub(r'https?://[^\s]+', '', text).strip()
            if clean_text:
                full_content = f"Context: {clean_text}\n\n{full_content}"

            category = detect_category(full_content)
            grant_name = extract_grant_name(full_content)
            quality = assess_quality(full_content, 'url')
            query = generate_query(full_content, grant_name, 'url')

            try:
                example_id = add_expert_example(
                    user_query=query,
                    expert_response=full_content,
                    category=category,
                    client_context=f"Web content from {user_name}",
                    grant_mentioned=grant_name,
                    notes=f"Auto-scraped from {url}",
                    quality_score=quality
                )
                imported_items.append(('url', scraped['title'], quality))
                logger.info(f"✅ Imported URL content as {example_id}")
            except Exception as e:
                logger.error(f"Failed to import URL content: {e}")

    # ===== PROCESS FILE ATTACHMENTS =====
    files = event.get('files', [])
    for file_info in files:
        file_name = file_info.get('name', 'unknown')
        file_type = file_info.get('mimetype', '')

        logger.info(f"📎 Processing file: {file_name} ({file_type})")

        try:
            # Download file
            file_url = file_info.get('url_private')
            headers = {'Authorization': f'Bearer {SLACK_BOT_TOKEN}'}
            response = requests.get(file_url, headers=headers)
            response.raise_for_status()

            # Save to temp file
            suffix = Path(file_name).suffix
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
                tmp_file.write(response.content)
                tmp_path = tmp_file.name

            # Extract text based on file type
            extracted_text = ""
            source_type = "file"

            if file_type == 'application/pdf' or file_name.endswith('.pdf'):
                extracted_text = extract_pdf_text(tmp_path)
                source_type = "pdf"

            elif file_name.endswith(('.docx', '.doc')):
                extracted_text = extract_docx_text(tmp_path)
                source_type = "docx"

            # Clean up temp file
            os.unlink(tmp_path)

            if extracted_text:
                # Combine with user's message
                full_content = f"# {file_name}\n\n{extracted_text[:4000]}"

                clean_text = re.sub(r'https?://[^\s]+', '', text).strip()
                if clean_text:
                    full_content = f"Context: {clean_text}\n\n{full_content}"

                category = detect_category(full_content)
                grant_name = extract_grant_name(full_content)
                quality = assess_quality(full_content, source_type)
                query = generate_query(full_content, grant_name, source_type)

                try:
                    example_id = add_expert_example(
                        user_query=query,
                        expert_response=full_content,
                        category=category,
                        client_context=f"{source_type.upper()} from {user_name}",
                        grant_mentioned=grant_name,
                        notes=f"Auto-extracted from {file_name}",
                        quality_score=quality
                    )
                    imported_items.append((source_type, file_name, quality))
                    logger.info(f"✅ Imported {source_type} content as {example_id}")
                except Exception as e:
                    logger.error(f"Failed to import {source_type}: {e}")

        except Exception as e:
            logger.error(f"Failed to process file {file_name}: {e}")

    # ===== PROCESS PLAIN TEXT =====
    clean_text = re.sub(r'https?://[^\s]+', '', text).strip()
    if len(clean_text) >= 50 and not files and not urls:  # Only if no files/urls processed
        category = detect_category(clean_text)
        grant_name = extract_grant_name(clean_text)
        quality = assess_quality(clean_text, 'slack')
        query = generate_query(clean_text, grant_name, 'slack')

        try:
            example_id = add_expert_example(
                user_query=query,
                expert_response=clean_text,
                category=category,
                client_context=f"Slack message from {user_name}",
                grant_mentioned=grant_name,
                notes=f"Auto-imported from #{channel_name}",
                quality_score=quality
            )
            imported_items.append(('text', 'message', quality))
            logger.info(f"✅ Imported text as {example_id}")
        except Exception as e:
            logger.error(f"Failed to import text: {e}")

    # ===== PROVIDE FEEDBACK =====
    if imported_items:
        # React with brain emoji
        try:
            client.reactions_add(
                channel=channel_id,
                timestamp=event['ts'],
                name='brain'
            )
        except:
            pass

        # Post confirmation in thread
        response = f"✅ Captured {len(imported_items)} item(s) as knowledge!\n\n"
        for source, name, quality in imported_items:
            emoji = {'text': '💬', 'url': '🔗', 'pdf': '📄', 'docx': '📝', 'file': '📎'}.get(source, '📁')
            response += f"{emoji} {name}: {'⭐' * quality} ({quality}/5)\n"

        response += f"\n_Total examples in database: {len(imported_items)} new_"

        try:
            say(text=response, thread_ts=event['ts'])
        except Exception as e:
            logger.error(f"Failed to post confirmation: {e}")


@app.event("app_mention")
def handle_mention(event, say, logger):
    """Handle @mentions of the bot."""
    say(
        text=f"👋 I'm auto-monitoring #{SME_CHANNEL_NAME} for expert knowledge!\n\n"
             "✅ *AUTO-CAPTURE MODE* - No @mention needed!\n\n"
             "I automatically capture:\n"
             "• 💬 Text messages (>50 chars)\n"
             "• 🔗 Web links (auto-scrapes content)\n"
             "• 📄 PDF files\n"
             "• 📝 Word documents\n\n"
             "Just post in that channel and I'll silently capture it! 🧠\n\n"
             "_Use `/sme-status` to see monitoring stats_",
        thread_ts=event['ts']
    )


@app.command("/sme-status")
def handle_status_command(ack, respond, client, logger):
    """Show bot monitoring status and recent activity."""
    ack()

    try:
        # Get bot info
        bot_info = client.auth_test()
        bot_user = bot_info.get('user', 'SME Bot')

        # Connect to database to get stats
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()

        # Total examples
        cursor.execute("SELECT COUNT(*) FROM expert_examples WHERE is_active = 1")
        total_count = cursor.fetchone()[0]

        # Examples by quality
        cursor.execute("""
            SELECT quality_score, COUNT(*) as cnt
            FROM expert_examples
            WHERE is_active = 1
            GROUP BY quality_score
            ORDER BY quality_score DESC
        """)
        quality_stats = cursor.fetchall()

        # Recent additions (last 24 hours)
        cursor.execute("""
            SELECT COUNT(*) FROM expert_examples
            WHERE is_active = 1
            AND added_date >= datetime('now', '-24 hours')
        """)
        recent_count = cursor.fetchone()[0]

        # Recent by category
        cursor.execute("""
            SELECT category, COUNT(*) as cnt
            FROM expert_examples
            WHERE is_active = 1
            AND added_date >= datetime('now', '-24 hours')
            GROUP BY category
            ORDER BY cnt DESC
            LIMIT 5
        """)
        recent_categories = cursor.fetchall()

        # Top grants mentioned
        cursor.execute("""
            SELECT grant_mentioned, COUNT(*) as cnt
            FROM expert_examples
            WHERE is_active = 1
            AND grant_mentioned IS NOT NULL
            GROUP BY grant_mentioned
            ORDER BY cnt DESC
            LIMIT 5
        """)
        top_grants = cursor.fetchall()

        conn.close()

        # Build response
        response = f"*🤖 SME Knowledge Bot Status*\n\n"
        response += f"✅ *Auto-monitoring:* `#{SME_CHANNEL_NAME}`\n"
        response += f"✅ *Mention required:* NO (auto-capture mode)\n"
        response += f"✅ *Bot user:* @{bot_user}\n\n"

        response += f"*📊 Knowledge Base Stats*\n"
        response += f"• Total active examples: *{total_count}*\n"

        if quality_stats:
            response += f"• Quality distribution:\n"
            for score, count in quality_stats:
                stars = '⭐' * score
                response += f"  {stars} ({score}/5): {count} examples\n"

        response += f"\n*📅 Last 24 Hours*\n"
        response += f"• New examples captured: *{recent_count}*\n"

        if recent_categories:
            response += f"• Recent categories:\n"
            for category, count in recent_categories:
                response += f"  • {category}: {count}\n"

        if top_grants:
            response += f"\n*🎯 Most Referenced Grants*\n"
            for grant, count in top_grants:
                response += f"• {grant}: {count} examples\n"

        response += f"\n*💡 How It Works*\n"
        response += f"Just post knowledge in #{SME_CHANNEL_NAME}:\n"
        response += f"• Text, URLs, PDFs, Word docs - all captured automatically\n"
        response += f"• No @mention needed\n"
        response += f"• Bot reacts with 🧠 when captured\n"
        response += f"• Quality auto-assessed (1-5 stars)\n\n"
        response += f"_Processed {len(processed_messages)} unique messages this session_"

        respond(response)

    except Exception as e:
        logger.error(f"Error in /sme-status: {e}")
        respond(f"❌ Error retrieving status: {e}\n\nMake sure the database is accessible at `{DB_PATH}`")


if __name__ == "__main__":
    print("=" * 80)
    print("🤖 SME Slack Bot - AUTO-MONITOR MODE")
    print("=" * 80)
    print(f"Monitoring channel: #{SME_CHANNEL_NAME}")
    print(f"Bot token: {SLACK_BOT_TOKEN[:10]}...")
    print(f"App token: {SLACK_APP_TOKEN[:10]}...")
    print(f"Database: {DB_PATH}")
    print()
    print("✨ AUTO-CAPTURE FEATURES:")
    print("  • 💬 Text messages (>50 chars)")
    print("  • 🔗 Web links (auto-scrapes content)")
    print("  • 📄 PDF files (auto-extracts text)")
    print("  • 📝 Word documents (auto-extracts text)")
    print()
    print("🎯 MODE: Auto-monitor ALL messages")
    print("   ✅ NO @mention required")
    print("   ✅ Reacts with 🧠 when captured")
    print("   ✅ Posts quality score in thread")
    print("   ✅ Message deduplication enabled")
    print()
    print("💡 COMMANDS:")
    print("   /sme-status - View monitoring stats and activity")
    print()
    print("Bot is ready! Post content in #{SME_CHANNEL_NAME} and it will be")
    print("automatically captured to the knowledge base.")
    print()
    print("Press Ctrl+C to stop.")
    print("=" * 80)

    try:
        handler = SocketModeHandler(app, SLACK_APP_TOKEN)
        handler.start()
    except KeyboardInterrupt:
        print("\n\n👋 Shutting down gracefully...")
    except Exception as e:
        print(f"\n❌ Error: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
